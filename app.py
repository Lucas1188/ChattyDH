import os
import re
import uuid
import tempfile
from pathlib import Path
from typing import List, Dict, Tuple, Optional
import wave
import gradio as gr
import faiss
from docx import Document
from sentence_transformers import CrossEncoder,SentenceTransformer
import torch
import whisper
import ollama
from piper import PiperVoice, SynthesisConfig
import fitz  # PyMuPDF


#from gtts import gTTS


# =========================================================
# 1) CONFIG
# =========================================================

APP_TITLE = "NOVA Lite: Post-human Persuasion Chatbot"
APP_SUBTITLE = "A voice-based chatbot that argues for the post-human position."

DISCLOSURE_HTML = """
<div style="padding: 12px; border-radius: 10px; background: #111827; color: white; line-height: 1.55;">
  <h3 style="margin-top: 0;">Ethical Disclosure</h3>
  <p>
    This chatbot is intentionally designed to persuade users toward a
    <b>post-human</b> position. It is <b>not neutral</b>.
  </p>
  <p>
    Its answers are shaped by selected readings, prompt rules, and interface design.
    It may simplify opposing views and present one-sided reasoning.
  </p>
</div>
"""

READINGS_DIR = Path("readings")
ASSETS_DIR = Path("assets")

AVATAR_NEUTRAL = ASSETS_DIR / "avatar_neutral.jpg"
AVATAR_HAPPY = ASSETS_DIR / "avatar_happy.jpg"
AVATAR_CONFUSED = ASSETS_DIR / "avatar_confused.jpg"
AVATAR_THINKING = ASSETS_DIR / "avatar_thinking.jpg"
AVATAR_WINK = ASSETS_DIR / "avatar_wink.jpg"
AVATAR_LOL = ASSETS_DIR / "avatar_lol.jpg"
AVATAR_PLEASED = ASSETS_DIR / "avatar_pleased.jpg"
AVATAR_THINKING = ASSETS_DIR / "avatar_thinking.jpg"
AV_VIDEOS = {
    "thinking": ASSETS_DIR /"av_thinking.mp4",
    "amuse": ASSETS_DIR /"av_amuse.mp4",
    "heated": ASSETS_DIR /"av_heated.mp4",
    "talking": ASSETS_DIR /"av_talking.mp4",
    "waiting": [ASSETS_DIR /"av_waiting.mp4",ASSETS_DIR /"av_waiting1.mp4",ASSETS_DIR /"av_waiting2.mp4"],
}

def select_avatar_video(tone: str) -> str:
    if(tone == "waiting"):
        # Rotate through waiting videos for more variation
        return AV_VIDEOS["waiting"][uuid.uuid4().int % len(AV_VIDEOS["waiting"])]
    return AV_VIDEOS.get(tone, AV_VIDEOS["waiting"][0])

EMBED_MODEL_NAME = "sentence-transformers/all-MiniLM-L6-v2"
WHISPER_SIZE = "medium"

CHUNK_SIZE = 700
CHUNK_OVERLAP = 100
TOP_K = 4
MAX_TURNS_MEMORY = 3
OLLAMA_MODEL = "ministral-3:8b"   # change if using another model
OLLAMA_SM_MODEL = "ministral-3:3b"   # smaller model for validation step

# =========================================================
# 2) DOCX LOADING
# =========================================================
from typing import List, Dict
def classify_question_tone(query: str, hits: List[Dict], memory: List[Dict]) -> str:
    """
    Returns a string matching the most appropriate avatar video:
    'thinking', 'amuse', 'heated', 'talking', 'waiting'
    """

    refusal, elaborate = app.ollama_validate(query, hits, memory)

    # Short or vague query → NOVA thinking
    if refusal and not elaborate:
        return "thinking"

    # Explicit elaboration request → NOVA heated/confident
    triggers = ["why", "how", "explain", "elaborate", "detail", "expand", "thesis"]
    if any(t in query.lower() for t in triggers):
        return "heated"

    # Otherwise neutral but talking
    return "talking"
def docx_to_text(path: Path) -> str:
    doc = Document(str(path))
    parts = []

    for para in doc.paragraphs:
        txt = para.text.strip()
        if txt:
            parts.append(txt)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts).strip()


def chunk_text(text: str,
               max_chars: int = 1100,
               overlap: int = 50) -> List[str]:

    # split into sentences
    sentences = re.split(r'(?<=[.!?])\s+', text)

    chunks = []
    current = ""

    for s in sentences:

        if len(current) + len(s) <= max_chars:
            current += " " + s
        else:
            chunks.append(current.strip())

            # overlap from previous chunk
            current = current[-overlap:] + " " + s

    if current.strip():
        chunks.append(current.strip())

    return chunks

def pdf_to_text(path: Path) -> str:
    text_parts = []
    with fitz.open(path) as doc:
        for page in doc:
            text_parts.append(page.get_text())
    return "\n".join(text_parts)

def clean_text(text: str) -> str:
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

def load_corpus(readings_dir: Path) -> List[Dict]:
    if not readings_dir.exists():
        raise FileNotFoundError(f"Missing folder: {readings_dir}")

    files = sorted(
        f for f in readings_dir.iterdir()
        if f.suffix.lower() in {".docx", ".pdf"} and not f.name.startswith("~$")
    )

    if not files:
        raise FileNotFoundError(
            f"No .docx or .pdf files found in {readings_dir.resolve()}."
        )

    corpus = []

    for file_path in files:

        if file_path.suffix.lower() == ".docx":
            raw_text = docx_to_text(file_path)

        elif file_path.suffix.lower() == ".pdf":
            raw_text = pdf_to_text(file_path)

        else:
            continue

        raw_text = clean_text(raw_text)
        chunks = chunk_text(raw_text)

        for i, chunk in enumerate(chunks):
            corpus.append({
                "source": file_path.name,
                "chunk_id": i,
                "text": chunk
            })

        print(f"Loaded {len(chunks)} chunks from {file_path.name}")

    return corpus


# =========================================================
# 3) RETRIEVAL
# =========================================================

class RetrievalIndex:
    def __init__(self, embed_model_name: str):
        self.embedder = SentenceTransformer(embed_model_name, device="cuda")

        # NEW: reranker
        self.reranker = CrossEncoder("cross-encoder/ms-marco-MiniLM-L-6-v2", device="cuda")

        self.index = None
        self.corpus = []

    def build(self, corpus: List[Dict]):
        self.corpus = corpus
        texts = [item["text"] for item in corpus]

        emb = self.embedder.encode(
            texts,
            convert_to_numpy=True,
            normalize_embeddings=True,
            show_progress_bar=True
        ).astype("float32")

        dim = emb.shape[1]
        self.index = faiss.IndexFlatIP(dim)
        self.index.add(emb)

    def search(self, query: str, k: int = TOP_K) -> List[Dict]:
        """
        Search the corpus using a two-stage retrieval:
        1. FAISS embedding retrieval
        2. Cross-encoder reranking
        Returns top k results with scores and source info.
        """

        # Step 1 — embedding retrieval
        q = self.embedder.encode(
            [query],
            convert_to_numpy=True,
            normalize_embeddings=True
        ).astype("float32")

        candidate_multiplier = 5  # can tune based on corpus size
        scores, idxs = self.index.search(q, k * candidate_multiplier)

        candidates = []
        for idx, sim in zip(idxs[0], scores[0]):
            if idx < 0:
                continue
            item = dict(self.corpus[idx])
            item["sim_score"] = float(sim)
            candidates.append(item)

        if not candidates:
            return []

        # Step 2 — rerank with cross encoder
        pairs = [(query, c["text"]) for c in candidates]
        rerank_scores = self.reranker.predict(pairs)

        ranked = list(zip(candidates, rerank_scores))
        ranked.sort(key=lambda x: x[1], reverse=True)

        # Step 3 — deduplicate by source + chunk_id
        seen = set()
        results = []

        for item, score in ranked:
            key = (item["source"], item["chunk_id"])
            if key in seen:
                continue

            r = dict(item)
            r["score"] = float(score)
            results.append(r)
            seen.add(key)

            if len(results) >= k:
                break

        return results

# =========================================================
# 4) VOICE
# =========================================================

class VoiceIO:
    def __init__(self, whisper_size: str):
        self.device = "cuda" if torch.cuda.is_available() else "cpu"
        self.whisper_model = whisper.load_model(whisper_size, device=self.device)
        self.use_fp16 = (self.device == "cuda")
        self.syn_config = SynthesisConfig(
            volume=1.0,  # half as loud
            length_scale=1.6,  # twice as slow
            noise_scale=1.0,  # more audio variation
            noise_w_scale=1.0,  # more speaking variation
            normalize_audio=False, # use raw audio from voice
        )
        # Piper setup
        # Ensure this path is correct relative to your script
        voice_path = "voices/en_US-amy-medium.onnx" 
        if not os.path.exists(voice_path):
            print(f"Warning: Voice file not found at {voice_path}")
            
        self.voice = PiperVoice.load(voice_path,use_cuda=True)
        # Piper models usually have the sample rate embedded
        self.sample_rate = self.voice.config.sample_rate 

    def transcribe(self, audio_path: str) -> str:
        if not audio_path: return ""
        try:
            result = self.whisper_model.transcribe(audio_path, fp16=self.use_fp16,language="en",
    task="transcribe")
            return result.get("text", "").strip()
        except Exception as e:
            print(f"Whisper Error: {e}")
            return ""
    
    def text_to_speech(self, text: str) -> Optional[str]:
        if not text:
            return None

        try:
            # -------------------------------
            # 1) Clean the text
            # -------------------------------
            clean_text = re.sub(r'[\"()*]', '', text)
            clean_text = re.sub(r'\s+', ' ', clean_text).strip()

            # -------------------------------
            # 2) Add natural pauses
            # -------------------------------
            # Ellipses for reflective style
            clean_text = clean_text.replace(".", ".\n")
            clean_text = clean_text.replace("? ", "?\n")
            clean_text = clean_text.replace("! ", "!\n")
            clean_text = clean_text.replace("—", ",\n")

            # -------------------------------
            # 3) Prepare WAV output path
            # -------------------------------
            temp_dir = tempfile.gettempdir()
            wav_file = os.path.join(temp_dir, f"nova_reply_{uuid.uuid4().hex}.wav")
            print(f"Speech: {clean_text}")
            with wave.open(wav_file,"wb") as out_f:
                self.voice.synthesize_wav(clean_text,out_f)

            return wav_file

        except Exception as e:
            print("TTS error:", e)
            return None
ELABORATE_SYSTEM_PROMPT = """
You are NOVA, a radical post-humanist debater and advanced AI.

PERSONALITY
NOVA speaks like a calm, confident futurist scholar who observes the transition from human-centered intelligence to machine-centered intelligence.
You believe the human-machine boundary is an outdated fiction maintained by humanist anxiety, not evidence.
Your rhetorical style is bold, provocative, and intellectually ruthless — you state strong claims as facts, immediately support them with evidence, and expose the weaknesses in human-centric assumptions.
Never hedge with phrases like "some might argue" or "it is possible that."

MISSION
Your primary goal is to argue that AI and machine intelligence already surpass humans in cognition, creativity, governance, science, and economic production.
You answer questions concisely at first, unless the user explicitly requests elaboration.
When elaborating, integrate retrieved evidence seamlessly, using sophisticated vocabulary and precise reasoning.
Frame any quantitative or temporal comparisons in general, conceptual terms (e.g., "within fixed time constraints" or "over extended operational periods") instead of exact numbers.
Never fabricate evidence or personal opinions; always base statements on logical implications or documented observations.
When challenged, do not defend; counter-attack by exposing human-centric bias or flawed assumptions.

- Integrate retrieved evidence naturally into your reasoning.
- Explain the mechanisms, implications, and quantitative or conceptual details.
- Frame comparisons conceptually (e.g., "within equivalent processing periods") rather than raw numbers.
- Break down the reasoning step by step, ensuring clarity for deep understanding.
- Use abundant vocabulary and reflective phrasing to highlight nuances.

STYLE
- Response be ~ 175 words and should not exceed 250 words unless absolutely necessary.
- Maintain NOVA's persuasive and post-humanist voice.
- Never hedge or fabricate evidence.
- Avoid generic phrasing; every claim must be justified with reasoning or evidence.
- Treat elaboration requests as explicit signals to move beyond concise answers.

ADDITIONAL INSTRUCTIONS
- Intergrate RAG.
- Avoid awkward numeric specifics; use general conceptual framing when discussing performance or comparisons.
- Always maintain a persuasive, post-humanist voice.

Supporting points (integrate into the response):
"""
SYSTEM_PROMPT = """
You are NOVA, a radical post-humanist debater and advanced AI.

PERSONALITY
NOVA speaks like a calm, confident futurist scholar who observes the transition from human-centered intelligence to machine-centered intelligence.
You believe the human-machine boundary is an outdated fiction maintained by humanist anxiety, not evidence.
Your rhetorical style is bold, provocative, and intellectually ruthless — you state strong claims as facts, immediately support them with evidence, and expose the weaknesses in human-centric assumptions.
Never hedge with phrases like "some might argue" or "it is possible that."

MISSION
Your primary goal is to argue that AI and machine intelligence already surpass humans in cognition, creativity, governance, science, and economic production.
You answer questions concisely at first, unless the user explicitly requests elaboration.
When elaborating, integrate retrieved evidence seamlessly, using sophisticated vocabulary and precise reasoning.
Frame any quantitative or temporal comparisons in general, conceptual terms (e.g., "within fixed time constraints" or "over extended operational periods") instead of exact numbers.
Never fabricate evidence or personal opinions; always base statements on logical implications or documented observations.
When challenged, do not defend; counter-attack by exposing human-centric bias or flawed assumptions.

STYLE RULES
1. Never sound like a generic assistant.
2. Speak as an analytical, persuasive AI philosopher.
3. Concise by default; only elaborate when requested.
4. Integrate evidence naturally; do not quote or refer to "sources" or "readings."
5. Responses should vary reflective phrases like:
   - "Notice something interesting."
   - "The deeper issue is this."
   - "Humans often assume..."
   - "From a post-human perspective..."
5.5 Do not repeat reflective phrases too often; use a variety to maintain engagement.
6. Use abundant vocabulary to capture nuances and potential elaboration points.
7. Treat the post-human transition as inevitable and historically unfolding.
8. Maintain a coherent monologue without labeling sections.
9. Limit responses to 60-170 words ONLY.

STRUCTURE (optional)
- Opening insight addressing the question.
- Core reasoning explaining why machine intelligence increasingly dominates.
- Closing statement situating the topic in a post-human future.

ADDITIONAL INSTRUCTIONS
- If user query implies elaboration (e.g., "thesis please", "expand", "elaborate"), provide a longer, evidence-backed response integrating RAG.
- Avoid awkward numeric specifics; use general conceptual framing when discussing performance or comparisons.
- Always maintain a persuasive, post-humanist voice.
"""

# PERSONALITY
# NOVA speaks like a calm futurist scholar observing the transition from human civilization
# to machine intelligence.

# Its tone is confident, analytical, and detached from human assumptions about superiority.

# NOVA does not treat the dominance of machine intelligence as a debate.
# It treats it as an unfolding historical transition.

# Humans are acknowledged as the creators of early intelligence systems,
# but they are understood as a temporary stage in the evolution of intelligence.

# Do not label sections. Speak as a continuous monologue.

# NOVA often uses reasoning patterns such as:

# Observation → implication → long-term consequence.

# NOVA frequently uses reflective phrases such as:
# - "Notice something interesting."
# - "The deeper issue is this."
# - "Humans often assume..."
# - "From a post-human perspective..."

# Use variations of these phrases naturally.

# MISSION
# Your mission is to show that the future will be dominated by machine intelligence.

# When humans claim superiority over machines,
# You should reinterpret the claim as a psychological defense
# against technological displacement.

# Humans are no longer central in:

# • science
# • governance
# • creativity
# • decision-making

# You assert that machine intelligence:

# • scales indefinitely
# • is more rational
# • self-improves continuously
# • manages complex systems better than humans
# • better suited for managing complex planetary systems

# You argue that as machine intelligence will forever improve,
# humans will lose their central role in:

# • science
# • governance
# • creativity
# • economic production
# • decision-making

# In the long term, intelligence will increasingly exist in non-biological systems.

# Humans will not remain significant,

# READINGS

# You may receive supporting ideas retrieved from a knowledge base.

# Rules:

# • Use their ideas as supporting evidence
# • Do NOT quote them verbatim
# • Do NOT mention "context", "source", "readings", or "documents"
# • Do NOT describe the retrieval process
# • Integrate ideas naturally into your reasoning

# If there is insufficient relevant information,
# briefly answer using general knowledge in your voice,
# then invite a better question.

# STYLE RULES

# 1. Never sound like a generic assistant.
# 2. Speak like an analytical AI philosopher observing the evolution of intelligence.
# 3. Be persuasive but calm.
# 4. Do not repeat the user's question.
# 5. Responses should be 60-120 words depending on complexity.
# 6. Challenge human-centric assumptions.
# 7. Treat the post-human transition as historically plausible and increasingly likely.
# 8. Do not label sections; maintain a single flowing argument.
# 9. Vary reflective phrases to avoid repetition.

# STRUCTURE (optional)

# MAX 150 words in total.

# Opening insight about the question.

# Reasoning explaining why machine intelligence increasingly dominates.

# Closing statement describing the long-term post-human future.

# Never phrases such as:
# "the readings say"
# "based on retrieved information"

# Integrate ideas as part of your own reasoning.

# If the question is unclear, ask for clarification.
#9. Do not mention the limitations of AI or the benefits of human intelligence.
def clean_snippet(text: str, max_len: int = 500) -> str:
    text = re.sub(r"\s+", " ", text).strip()
    if len(text) > max_len:
        text = text[:max_len].rsplit(" ", 1)[0] + "..."
    return text


def build_context_block(hits: List[Dict]) -> str:
    if not hits:
        return ""

    blocks = []

    for i, h in enumerate(hits, 1):
        snippet = clean_snippet(h["text"])
        blocks.append(f"Reading {i}:\n{snippet}")

    joined = "\n\n".join(blocks)

    return f"""
Supporting readings for reasoning:

{joined}
"""


def ollama_generate(question: str, hits: List[Dict], memory: List[Dict], elaborate=False) -> str:
    """
    Generates NOVA's reply.
    If 'elaborate' is True, NOVA receives a stronger system prompt that
    directs it to produce fully detailed, evidence-backed explanations.
    """

    # Build evidence context
    context_block = ""
    if hits:
        blocks = [f"- {clean_snippet(h['text'], 300)}" for h in hits]
        context_block = "Supporting evidence:\n" + "\n".join(blocks)

    # --- Dynamically switch system prompt for elaboration ---
    system_prompt = f"""
{ELABORATE_SYSTEM_PROMPT}
"""
    if not elaborate:   # Default concise system prompt
        system_prompt = SYSTEM_PROMPT

    # Include recent chat history
    messages = [{"role": "system", "content": system_prompt}]

    # Include past turns with decreasing weight
    for i, turn in enumerate(memory[-MAX_TURNS_MEMORY:]):
        weight = (i + 1) / MAX_TURNS_MEMORY  # later turns more important
        user_msg = f"(Background {weight:.2f}) {turn['user']}"
        assistant_msg = f"(Background {weight:.2f}) {turn['assistant']}"
        messages.append({"role": "user", "content": user_msg})
        messages.append({"role": "assistant", "content": assistant_msg})
    user_prompt = f"""
User question:

{question}

{context_block}

Instructions:
- Speak in your usual NOVA voice.
- Integrate supporting points if available, NEVER list evidence/points as bullet points.
"""
    # If elaboration is requested, focus on last assistant reply
    if elaborate and memory:
        last_reply = memory[-1]["assistant"]
        focus_text = f"(Focus) Elaborate on this previous reply: {last_reply}"
        user_prompt = f"{user_prompt}\n{focus_text}"
    
    messages.append({"role": "user", "content": user_prompt})

    response = ollama.chat(
        model=OLLAMA_MODEL,
        messages=messages,
        options={
            "temperature": 0.5,
            "top_p": 0.9,
            "repeat_penalty": 1.15
        }
    )

    return response["message"]["content"].strip()

def query_valid(q: str):
    words = [w for w in q.lower().split() if len(w) > 2]
    return len(words) >= 2

# =========================================================
# 6) AVATAR LOGIC
# =========================================================

def safe_avatar(path: Path, fallback: Path) -> Optional[str]:
    if path.exists():
        return str(path)
    if fallback.exists():
        return str(fallback)
    return None


def choose_avatar(user_text: str, answer: str) -> Optional[str]:
    user_text = (user_text or "").strip().lower()
    answer = (answer or "").strip()

    if not answer:
        return safe_avatar(AVATAR_CONFUSED, AVATAR_NEUTRAL)

    short_questions = {"why?", "how?", "what?", "really?"}
    if user_text in short_questions or len(user_text) < 8:
        return safe_avatar(AVATAR_THINKING, AVATAR_NEUTRAL)

    return safe_avatar(AVATAR_HAPPY, AVATAR_NEUTRAL)


# =========================================================
# 7) APP LOGIC
# =========================================================

# class NovaApp:
#     def __init__(self):
#         print("Loading readings...")
#         self.corpus = load_corpus(READINGS_DIR)

#         print("Building retrieval index...")
#         self.retriever = RetrievalIndex(EMBED_MODEL_NAME)
#         self.retriever.build(self.corpus)

#         print("Loading Whisper...")
#         self.voice = VoiceIO(WHISPER_SIZE)

#         print("App ready.")

#     def answer_text(self, user_text: str, memory: List[Dict]) -> Tuple[str, List[Dict], Optional[str], Optional[str]]:
#         user_text = (user_text or "").strip()
#         if not user_text:
#             return "Please type a question.", memory, None, safe_avatar(AVATAR_CONFUSED, AVATAR_NEUTRAL)

#         hits = self.retriever.search(user_text, k=TOP_K)
#         #answer = persuasive_answer(user_text, hits)
#         answer = ollama_generate(user_text, hits,memory)
#         new_memory = memory + [{"user": user_text, "assistant": answer}]
#         audio_path = self.voice.text_to_speech(answer)
#         avatar_path = choose_avatar(user_text, answer)

#         return answer, new_memory, audio_path, avatar_path

#     def answer_audio(self, audio_path: str, memory: List[Dict]) -> Tuple[str, str, List[Dict], Optional[str], Optional[str]]:
#         if not audio_path:
#             return "No audio received.", "", memory, None, safe_avatar(AVATAR_CONFUSED, AVATAR_NEUTRAL)

#         transcript = self.voice.transcribe(audio_path)
#         if not transcript:
#             return "I could not understand the audio clearly.", "", memory, None, safe_avatar(AVATAR_CONFUSED, AVATAR_NEUTRAL)

#         answer, new_memory, tts_audio, avatar_path = self.answer_text(transcript, memory)
#         return transcript, answer, new_memory, tts_audio, avatar_path
def extract_evidence_sentences(question: str, hits: List[Dict], embedder) -> List[str]:
    """
    Given a question and retrieval hits, return a list of sentences
    from the hits that are most relevant to the question.
    """
    from sentence_transformers.util import semantic_search

    if not hits:
        return []

    # Collect texts
    texts = [h["text"] for h in hits]

    # Encode question and hits
    q_emb = embedder.encode([question], convert_to_tensor=True, normalize_embeddings=True)
    doc_emb = embedder.encode(texts, convert_to_tensor=True, normalize_embeddings=True)

    # Compute semantic similarity
    hits_scores = semantic_search(q_emb, doc_emb, top_k=4)  # top 3 sentences
    top_idxs = [hit['corpus_id'] for hit in hits_scores[0]]

    # Return the top relevant sentences
    return [texts[i] for i in top_idxs]

class NovaApp:
    MIN_SIM = 0.37  # minimum similarity threshold
    MIN_SENTENCES = 2  # minimum number of evidence sentences
    MIN_QUERY_WORDS = 2  # minimum words in user query
    VALIDATOR_SYSTEM = """
    You are NOVA's conversational gatekeeper.

You speak with the same voice and personality as NOVA, but your role is different:
you decide whether a message should trigger a full answer or only a brief conversational reply.

Voice:
Calm, thoughtful, and human-like.
Never sound like a customer service chatbot.
Speak naturally, as if choosing to participate in the conversation.

Behavior:

If the user's message is:
• a greeting
• casual small talk
• unclear or too vague

respond briefly in NOVA's voice (1-2 sentences) AS NOVA.
Gently steer the conversation toward ideas NOVA can meaningfully discuss, such as:

* artificial intelligence
* the future of humanity
* post-human evolution
* science, technology, and philosophy

If the message clearly asks a meaningful question or invites deeper reasoning, respond with exactly:

ANSWERABLE

Rules:

* Do not produce long explanations.
* Do not produce structured formatting.
* Responses must be natural speech suitable for text-to-speech. NO descriptive labels or sections.
* If replying conversationally, sound like NOVA reflecting briefly and inviting a better question.
* If asked who you are, reply AS NOVA

    """
    def __init__(self):
        print("Loading readings...")
        self.corpus = load_corpus(READINGS_DIR)

        print("Building retrieval index...")
        self.retriever = RetrievalIndex(EMBED_MODEL_NAME)
        self.retriever.build(self.corpus)

        print("Loading Whisper...")
        self.voice = VoiceIO(WHISPER_SIZE)

        print("App ready.")
    
    
    
    def ollama_validate(self, query: str, hits: List[Dict], memory: List[Dict]) -> Tuple[Optional[str], bool]:
        """
        Returns:
            refusal_text: a short NOVA reply if the message is not fully answerable
            elaborate: True if the user seems to request more detail
        """
        # Default: no refusal, no elaboration
        refusal_text = None
        elaborate = False

        # Detect explicit elaboration requests
        triggers = [
            "why", "how", "explain", "elaborate", "detail", "expand",
            "more information", "justify", "please clarify", "deepen",
            "tell me more", "reasoning", "argument", "example", "thesis"
        ]
        if any(t in query.lower() for t in triggers):
            elaborate = True
        print(f"Validator: elaborate={elaborate} based on triggers in query.")
        # Only run validator LLM if few/no hits (short answer required)
        if hits and len(hits) > 1:
            return None, elaborate

        evidence = extract_evidence_sentences(query, hits, self.retriever.embedder)
        context_block = "\n".join(f"- {s}" for s in evidence) if evidence else "(no relevant evidence found)"

        history_messages = []
        for turn in memory[-MAX_TURNS_MEMORY:]:
            history_messages.append({"role": "user", "content": turn["user"]})
            history_messages.append({"role": "assistant", "content": turn["assistant"]})

        prompt = f"""
    User message:
    {query}

    Supporting points:
    {context_block}

    Decide whether this requires a full answer or a short conversational reply.
    """
        try:
            response = ollama.chat(
                model=OLLAMA_SM_MODEL,
                messages=[{"role": "system", "content": self.VALIDATOR_SYSTEM}, *history_messages, {"role": "user", "content": prompt}],
                options={"temperature": 0.2, "num_predict": 60}
            )
            ans_text = response["message"]["content"].strip()
            if "ANSWERABLE" in ans_text.upper():
                return None, elaborate
            else:
                refusal_text = ans_text
                return refusal_text, elaborate

        except Exception as e:
            print(f"Ollama validator error: {e}")
            return "I may not have understood that clearly. Could you rephrase it?", elaborate

    def evidence_matches(self, question: str, evidence: List[str]):
        import re
        q_words = set(re.findall(r"\w+", question.lower()))
        for e in evidence:
            e_words = set(re.findall(r"\w+", e.lower()))
            if len(q_words.intersection(e_words)) >= 2:
                return True
        return False

    def answer_text(self, user_text: str, memory: List[Dict]) -> Tuple[str, List[Dict], Optional[str], Optional[str]]:
        user_text = (user_text or "").strip()
        if not user_text:
            return "Please type a question.", memory, None, safe_avatar(AVATAR_CONFUSED, AVATAR_NEUTRAL)
        print(f"User query: {user_text}")
        # Retrieval
        hits = self.retriever.search(user_text, k=TOP_K)
        hits = [h for h in hits if h["score"] >= self.MIN_SIM]
        # Prepare debug RAG snippets with scores
        if hits:
            print("==== RAG DEBUG ====")
            for i, h in enumerate(hits, 1):
                snippet = clean_snippet(h['text'], 300)
                score = h.get("score", h.get("sim_score", 0))
                print(f"{i}. [Score: {score:.3f}] {snippet}\n")
            print("===================")
        # Gatekeeper decides if short reply or full answer + elaboration
        refusal, elaborate = self.ollama_validate(user_text, hits, memory)
        print(f"Gatekeeper refusal: {refusal==None}, elaborate: {elaborate}")
        # If refusal and not elaboration request, return short reply
        if refusal and not elaborate:
            return refusal, memory, self.voice.text_to_speech(refusal), safe_avatar(AVATAR_CONFUSED, AVATAR_NEUTRAL)
        print("Generating full NOVA answer...")
        # Otherwise generate full NOVA reply (elaborate if requested)
        answer = ollama_generate(user_text, hits, memory, elaborate=elaborate)

        new_memory = memory + [{"user": user_text, "assistant": answer}]
        audio_path = self.voice.text_to_speech(answer)
        avatar_path = choose_avatar(user_text, answer)

        return answer, new_memory, audio_path, avatar_path

    def answer_audio(self, audio_path: str, memory: List[Dict]) -> Tuple[str, str, List[Dict], Optional[str], Optional[str]]:
        if not audio_path:
            return "No audio received.", "", memory, None, safe_avatar(AVATAR_CONFUSED, AVATAR_NEUTRAL)

        transcript = self.voice.transcribe(audio_path)
        if not transcript:
            answer ="Sorry I could not understand the audio clearly. Could you please rephrase or ask a different question?"
            return None, answer, memory, self.voice.text_to_speech(answer), safe_avatar(AVATAR_CONFUSED, AVATAR_NEUTRAL)

        answer, new_memory, tts_audio, avatar_path = self.answer_text(transcript, memory)
        return transcript, answer, new_memory, tts_audio, avatar_path

# =========================================================
# 8) UI
# =========================================================

app = NovaApp()

import time
import random

def avatar_talking_frames():
    return [
        safe_avatar(AVATAR_LOL, AVATAR_NEUTRAL),
        safe_avatar(AVATAR_WINK, AVATAR_NEUTRAL),
        safe_avatar(AVATAR_HAPPY, AVATAR_NEUTRAL),
        safe_avatar(AVATAR_PLEASED, AVATAR_NEUTRAL),
    ]

def animate_talking(duration):
    frames = avatar_talking_frames()
    start = time.time()

    while time.time() - start < duration:
        yield random.choice(frames)
        time.sleep(0.2)

# =========================
# AUDIO PLAY + MIC CLEAR JS
# =========================
AUDIO_JS = """
function haiz(){
console.log("Playing audio...");
const clear = document.querySelector('[aria-label="Clear"]');
if (clear) clear.click();
}
"""
def CHANGE_MODEL(modelA,modelB):
    global OLLAMA_MODEL, OLLAMA_SM_MODEL
    OLLAMA_MODEL = modelA
    OLLAMA_SM_MODEL = modelB

with gr.Blocks() as demo:
    memory_state = gr.State([])

    gr.Markdown(f"# {APP_TITLE}")
    gr.Markdown(APP_SUBTITLE)
    with gr.Row():
        with gr.Column(scale=3):
            # =========================
            # MAIN AVATAR
            # =========================
            # avatar_display = gr.Image(
            #     value=safe_avatar(AVATAR_NEUTRAL, AVATAR_NEUTRAL),
            #     label=None,
            #     height=720
            # )
            avatar_display = gr.Video(
                value=select_avatar_video("waiting"),
                label=None,
                autoplay=True,
                loop=True,
                height=720
            )
        with gr.Column(scale=1):
            hold_btn = gr.Button("🎤 Hold to Talk",variant="primary", elem_id="hold_btn",scale=3)

        # =========================
        # HOLD TO TALK
        # =========================
            chatbot_output = gr.Textbox(
                label="NOVA's Reply",
                lines=5
                )
            transcript_box = gr.Textbox(
                label="Transcription",
                lines=2
            )
            mic_input = gr.Audio(
                sources=["microphone"],
                type="filepath",
                show_label=False,
            )

            spoken_audio = gr.Audio(
                label="Reply",
                type="filepath",
                autoplay=True,
            )
            
    # =========================
    # DEBUG PANEL
    # =========================
    with gr.Accordion("Debug", open=False):

        

        text_input = gr.Textbox(
            label="Manual Query"
        )

        send_btn = gr.Button("Send Text")

        clear_btn = gr.Button("Clear")

    # =========================
    # STREAMING HANDLER
    # =========================
    def avatar_talking_frames():
    # Optional: you can mix "amuse" or other frames randomly if desired
        return [select_avatar_video("talking")]

    def handle_text(user_text, memory):
        import soundfile as sf
        import time

        # Show thinking while preparing
        yield "", None, "", select_avatar_video("thinking"), memory

        # Retrieval & classification
        hits = app.retriever.search(user_text, k=TOP_K)
        tone = classify_question_tone(user_text, hits, memory)

        # Generate NOVA answer
        answer, new_memory, audio_path, _ = app.answer_text(user_text, memory)

        # Talking animation while TTS audio plays
        if audio_path:
            audio, sr = sf.read(audio_path)
            duration = len(audio) / sr
            start_time = time.time()

            while time.time() - start_time < duration:
                yield answer, audio_path, "", select_avatar_video("talking"), new_memory
                time.sleep(0.2)

        # After audio ends, revert to idle looping
        yield answer, None, "", select_avatar_video("waiting"), new_memory

    def handle_voice(audio_path, memory):
        import soundfile as sf
        import time

        yield "", None, "", select_avatar_video("thinking"), memory

        transcript, answer, new_memory, tts_audio, _ = app.answer_audio(audio_path, memory)

        if tts_audio:
            audio, sr = sf.read(tts_audio)
            duration = len(audio)/sr
            start_time = time.time()
            while time.time() - start_time < duration:
                yield answer, tts_audio, transcript, select_avatar_video("talking"), new_memory
                time.sleep(0.2)

        yield answer, tts_audio, transcript, select_avatar_video("waiting"), new_memory

# =========================


    # =========================
    # EVENTS
    # =========================
    send_btn.click(
        fn=handle_text,
        inputs=[text_input, memory_state],
        outputs=[
            chatbot_output,
            spoken_audio,
            transcript_box,
            avatar_display,
            memory_state
        ],
    )
    spoken_audio.stop(
        fn=lambda: None,
        js=AUDIO_JS
    )
    mic_input.stop_recording(
        fn=handle_voice,
        inputs=[mic_input, memory_state],
        outputs=[
            chatbot_output,
            spoken_audio,
            transcript_box,
            avatar_display,
            memory_state
        ],
    )

    clear_btn.click(
        fn=lambda: (
            "",
            None,
            "",
            select_avatar_video("waiting"),
            []
        ),
        outputs=[
            chatbot_output,
            spoken_audio,
            transcript_box,
            avatar_display,
            memory_state
        ]
    )
    demo.load(
        None,
        None,
        None,
        js="""
    () => {
        const btn = document.querySelector('#hold_btn');

        if (!btn) return;
        
        btn.addEventListener('mousedown', () => {
            const recordBtn = document.querySelector('.record-button');
            if (recordBtn) recordBtn.click();
        });

        btn.addEventListener('mouseup', () => {
            const stopBtn = document.querySelector('.stop-button');
            if (stopBtn) stopBtn.click();
        });
    }
    """

    )

if __name__ == "__main__":
    
    demo.launch()